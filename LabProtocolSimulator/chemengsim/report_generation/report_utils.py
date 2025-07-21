"""
Report Generation Utilities
==========================

Helper functions for creating reports from lab observation data.
"""

import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import io
import base64
from docxtpl import DocxTemplate
from docx import Document
import streamlit as st
import plotly.graph_objects as go
import plotly.express as px
from datetime import datetime
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import shutil

# Define paths
TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates")
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")

# Ensure directories exist
os.makedirs(TEMPLATE_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

def fig_to_base64(fig):
    """Convert a matplotlib figure to base64 encoded image"""
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=300, bbox_inches='tight')
    buf.seek(0)
    img_str = base64.b64encode(buf.read()).decode()
    return img_str

def create_downloadable_docx(document, filename):
    """Create a downloadable link for a docx document"""
    # Save document to temp file
    temp_file = os.path.join(OUTPUT_DIR, filename)
    document.save(temp_file)
    
    # Read the saved file and create download link
    with open(temp_file, "rb") as file:
        bytes_data = file.read()
    
    # Create download button
    st.download_button(
        label="Download Report",
        data=bytes_data,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

def create_table_dataframe(column_names, num_rows=1):
    """Create an empty dataframe with the given column names"""
    df = pd.DataFrame(columns=column_names)
    for _ in range(num_rows):
        df = pd.concat([df, pd.DataFrame([[""] * len(column_names)], columns=column_names)], ignore_index=True)
    return df

def insert_plotly_figure(fig, document, width=6, height=4):
    """Save a plotly figure as image and insert into document"""
    # Save plotly figure as png
    img_path = os.path.join(OUTPUT_DIR, "temp_plot.png")
    fig.write_image(img_path, width=width*100, height=height*100)
    
    # Insert the image
    from docx.shared import Inches
    document.add_picture(img_path, width=Inches(width))
    
    # Clean up
    if os.path.exists(img_path):
        os.remove(img_path)

def create_matplotlib_graph(x_data, y_data, title, xlabel, ylabel, legend=None):
    """Create a matplotlib graph from data"""
    fig, ax = plt.subplots(figsize=(10, 6))
    
    if isinstance(y_data[0], list):
        for i, y in enumerate(y_data):
            ax.plot(x_data, y, marker='o', label=legend[i] if legend else f"Series {i+1}")
    else:
        ax.plot(x_data, y_data, marker='o')
    
    ax.set_title(title)
    ax.set_xlabel(xlabel)
    ax.set_ylabel(ylabel)
    ax.grid(True)
    
    if legend:
        ax.legend()
    
    return fig

def create_plotly_graph(x_data, y_data, title, xlabel, ylabel, legend=None):
    """Create a plotly graph from data"""
    fig = go.Figure()
    
    if isinstance(y_data[0], list):
        for i, y in enumerate(y_data):
            fig.add_trace(go.Scatter(
                x=x_data, 
                y=y, 
                mode='lines+markers',
                name=legend[i] if legend else f"Series {i+1}"
            ))
    else:
        fig.add_trace(go.Scatter(
            x=x_data, 
            y=y_data, 
            mode='lines+markers'
        ))
    
    fig.update_layout(
        title=title,
        xaxis_title=xlabel,
        yaxis_title=ylabel,
        template="plotly_white"
    )
    
    return fig

def save_input_data(experiment_name, data):
    """Save user input data to a file"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{experiment_name}_{timestamp}.json"
    file_path = os.path.join(OUTPUT_DIR, filename)
    
    # Convert to JSON and save
    data.to_json(file_path)
    return filename

def add_experiment_metadata(document, title, student_name, student_id, date=None):
    """Add experiment metadata to document"""
    document.add_heading(title, level=1)
    
    # Add student info
    document.add_paragraph(f"Student Name: {student_name}")
    document.add_paragraph(f"Student ID: {student_id}")
    
    # Add date
    if date is None:
        date = datetime.now().strftime("%B %d, %Y")
    document.add_paragraph(f"Date: {date}")
    
    # Add separator
    document.add_paragraph("_" * 50)

def create_download_link(file_content, file_name, text="Download File"):
    """
    Creates a download link for any file content.
    
    Parameters:
    -----------
    file_content : bytes
        Binary content of the file to be downloaded
    file_name : str
        Name of the file to be downloaded
    text : str
        Text to display for the download link
        
    Returns:
    --------
    str
        HTML string containing the download link
    """
    b64 = base64.b64encode(file_content).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{file_name}">{text}</a>'
    return href

def generate_report_download_link(document, filename=None):
    """
    Generates a download link for a Word document.
    
    Parameters:
    -----------
    document : Document
        The python-docx Document object
    filename : str, optional
        Name of the file to be downloaded. If None, a default name is used.
        
    Returns:
    --------
    str
        HTML string containing the download link
    """
    if filename is None:
        filename = f"lab_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
    # Save document to a bytes stream
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    
    # Create download link
    file_bytes = file_stream.read()
    return create_download_link(file_bytes, filename, "Download Report (DOCX)")