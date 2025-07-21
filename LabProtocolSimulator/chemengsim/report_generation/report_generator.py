"""
Report Generator Module
======================

Primary module for generating reports from lab observation data.
"""

import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import io
import base64
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import datetime
from . import report_utils

class ReportGenerator:
    """Base class for generating experiment reports"""
    
    def __init__(self, experiment_name, experiment_title):
        self.experiment_name = experiment_name
        self.experiment_title = experiment_title
        self.document = None
        self.data = {}
        
    def create_new_document(self):
        """Create a new document for the report"""
        self.document = Document()
        
    def add_title_page(self, student_info):
        """Add a title page to the document"""
        # Add title
        title = self.document.add_heading(self.experiment_title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_para = self.document.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(f"Date: {datetime.datetime.now().strftime('%B %d, %Y')}").bold = True
        
        # Add student information
        if isinstance(student_info, list):
            students_para = self.document.add_paragraph()
            students_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            students_para.add_run("Submitted by:").bold = True
            
            for student in student_info:
                student_para = self.document.add_paragraph()
                student_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                student_para.add_run(f"{student['name']} ({student['id']})").italic = True
        else:
            student_para = self.document.add_paragraph()
            student_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            student_para.add_run("Submitted by:").bold = True
            
            student_info_para = self.document.add_paragraph()
            student_info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            student_info_para.add_run(f"{student_info['name']} ({student_info['id']})").italic = True
            
        # Add page break
        self.document.add_page_break()
        
    def add_aim_and_objective(self, aim, objectives):
        """Add aim and objectives to the document"""
        # Add aim
        self.document.add_heading("Aim", level=1)
        self.document.add_paragraph(aim)
        
        # Add objectives
        self.document.add_heading("Objectives", level=1)
        obj_para = self.document.add_paragraph()
        
        if isinstance(objectives, list):
            for idx, objective in enumerate(objectives, 1):
                obj_para.add_run(f"{idx}. {objective}\n")
        else:
            obj_para.add_run(objectives)
            
    def add_theory(self, theory_text):
        """Add theory section to the document"""
        self.document.add_heading("Theory", level=1)
        
        if isinstance(theory_text, list):
            for paragraph in theory_text:
                self.document.add_paragraph(paragraph)
        else:
            self.document.add_paragraph(theory_text)
            
    def add_observation_table(self, table_data, table_title="Observations"):
        """Add observation data table to the document"""
        self.document.add_heading(table_title, level=1)
        
        # Create table
        if isinstance(table_data, pd.DataFrame):
            # Add table header
            table = self.document.add_table(rows=1, cols=len(table_data.columns))
            table.style = "Table Grid"
            header_cells = table.rows[0].cells
            
            for i, column in enumerate(table_data.columns):
                header_cells[i].text = column
                header_cells[i].paragraphs[0].runs[0].bold = True
                
            # Add table data
            for _, row in table_data.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)
        else:
            st.error("Table data must be a pandas DataFrame")
            
    def add_calculations(self, calculations_text, sample_calculations=None):
        """Add calculations section to the document"""
        self.document.add_heading("Calculations", level=1)
        
        if sample_calculations:
            self.document.add_heading("Sample Calculations", level=2)
            
            if isinstance(sample_calculations, list):
                for calc in sample_calculations:
                    self.document.add_paragraph(calc)
            else:
                self.document.add_paragraph(sample_calculations)
        
        if isinstance(calculations_text, list):
            for paragraph in calculations_text:
                self.document.add_paragraph(paragraph)
        else:
            self.document.add_paragraph(calculations_text)
            
    def add_results_table(self, results_data, table_title="Results"):
        """Add results data table to the document"""
        self.document.add_heading(table_title, level=1)
        
        # Create table
        if isinstance(results_data, pd.DataFrame):
            # Add table header
            table = self.document.add_table(rows=1, cols=len(results_data.columns))
            table.style = "Table Grid"
            header_cells = table.rows[0].cells
            
            for i, column in enumerate(results_data.columns):
                header_cells[i].text = column
                header_cells[i].paragraphs[0].runs[0].bold = True
                
            # Add table data
            for _, row in results_data.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)
        else:
            st.error("Results data must be a pandas DataFrame")
            
    def add_graph(self, fig, caption=None, width=6):
        """Add a matplotlib figure to the document"""
        # Save figure to memory
        img_stream = io.BytesIO()
        fig.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
        img_stream.seek(0)
        
        # Add figure to document
        self.document.add_picture(img_stream, width=Inches(width))
        
        # Add caption if provided
        if caption:
            caption_para = self.document.add_paragraph(caption)
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_para.style = 'Caption'
            
        plt.close(fig)
        
    def add_discussion(self, discussion_text):
        """Add discussion section to the document"""
        self.document.add_heading("Discussion", level=1)
        
        if isinstance(discussion_text, list):
            for paragraph in discussion_text:
                self.document.add_paragraph(paragraph)
        else:
            self.document.add_paragraph(discussion_text)
            
    def add_conclusion(self, conclusion_text):
        """Add conclusion section to the document"""
        self.document.add_heading("Conclusion", level=1)
        
        if isinstance(conclusion_text, list):
            for paragraph in conclusion_text:
                self.document.add_paragraph(paragraph)
        else:
            self.document.add_paragraph(conclusion_text)
            
    def save_document(self, filename=None):
        """Save the document to a file"""
        if filename is None:
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{self.experiment_name}_report_{timestamp}.docx"
            
        file_path = os.path.join(report_utils.OUTPUT_DIR, filename)
        self.document.save(file_path)
        return file_path
        
    def create_downloadable_report(self):
        """Create a downloadable report"""
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{self.experiment_name}_report_{timestamp}.docx"
        
        # Generate download link
        download_link = report_utils.generate_report_download_link(self.document, filename)
        
        # Display download link
        st.markdown(download_link, unsafe_allow_html=True)
        
        # Also provide the standard Streamlit download option
        # Save to a bytes buffer
        buffer = io.BytesIO()
        self.document.save(buffer)
        buffer.seek(0)
        
        # Offer download via Streamlit
        st.download_button(
            label="⬇️ Download Lab Report",
            data=buffer,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

def display_report_generation_ui():
    """Display the main report generation UI"""
    st.title("Report Generation Module")
    
    st.markdown("""
    This module allows you to generate professional lab reports from your observation data. 
    Select an experiment from the sidebar, enter your observations, and the application will:
    
    1. Perform all the necessary calculations
    2. Generate graphs and visual representations
    3. Create a complete, formatted lab report
    4. Provide a downloadable Word document
    """)
    
    st.info("Select an experiment from the sidebar to begin creating your report.")